"""
Utility Functions for Modern RAG System
=======================================
Helper functions for document processing, validation, and logging.
"""

import logging
import sys
from typing import Union, List, Optional
from io import BytesIO
from pathlib import Path

import streamlit as st
from pypdf import PdfReader
from docx import Document
from langchain_community.document_loaders import WebBaseLoader
from langchain_core.documents import Document as LangChainDocument


def setup_logging(log_level: str = "INFO", log_file: Optional[str] = None) -> logging.Logger:
    """
    Configure logging for the application.
    
    Args:
        log_level: Logging level (DEBUG, INFO, WARNING, ERROR, CRITICAL)
        log_file: Optional log file path
        
    Returns:
        Configured logger instance
    """
    logger = logging.getLogger('RAGSystem')
    logger.setLevel(getattr(logging, log_level.upper()))
    
    # Clear existing handlers
    logger.handlers.clear()
    
    # Console handler with formatting
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(getattr(logging, log_level.upper()))
    
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)
    
    # File handler if specified
    if log_file:
        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(getattr(logging, log_level.upper()))
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
    
    return logger


def validate_input(input_type: str, input_data) -> bool:
    """
    Validate input data based on type.
    
    Args:
        input_type: Type of input (Link, PDF, Text, DOCX, TXT)
        input_data: The input data to validate
        
    Returns:
        True if valid, False otherwise
    """
    if not input_data:
        return False
    
    if input_type == "Link":
        if isinstance(input_data, list):
            # Filter out empty strings and validate URLs
            valid_urls = [url for url in input_data if url and url.strip()]
            return len(valid_urls) > 0
        return False
    
    elif input_type == "Text":
        return isinstance(input_data, str) and len(input_data.strip()) > 0
    
    elif input_type in ["PDF", "DOCX", "TXT"]:
        # Check if it's a valid file upload
        return input_data is not None
    
    return False


def process_pdf(input_data: Union[BytesIO, st.runtime.uploaded_file_manager.UploadedFile]) -> str:
    """
    Extract text from PDF file.
    
    Args:
        input_data: PDF file as BytesIO or UploadedFile
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If PDF processing fails
    """
    try:
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif hasattr(input_data, 'read'):  # UploadedFile
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data format for PDF")
        
        text = ""
        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        
        if not text.strip():
            raise ValueError("No text content could be extracted from the PDF")
        
        return text.strip()
    
    except Exception as e:
        raise ValueError(f"Error processing PDF: {str(e)}")


def process_docx(input_data: Union[BytesIO, st.runtime.uploaded_file_manager.UploadedFile]) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        input_data: DOCX file as BytesIO or UploadedFile
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If DOCX processing fails
    """
    try:
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif hasattr(input_data, 'read'):  # UploadedFile
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data format for DOCX")
        
        text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        if not text.strip():
            raise ValueError("No text content could be extracted from the DOCX file")
        
        return text.strip()
    
    except Exception as e:
        raise ValueError(f"Error processing DOCX: {str(e)}")


def process_txt(input_data: Union[BytesIO, st.runtime.uploaded_file_manager.UploadedFile]) -> str:
    """
    Extract text from TXT file.
    
    Args:
        input_data: TXT file as BytesIO or UploadedFile
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If TXT processing fails
    """
    try:
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif hasattr(input_data, 'read'):  # UploadedFile
            text = input_data.read().decode('utf-8')
        else:
            raise ValueError("Invalid input data format for TXT")
        
        if not text.strip():
            raise ValueError("The text file is empty")
        
        return text.strip()
    
    except UnicodeDecodeError:
        raise ValueError("File encoding is not UTF-8. Please save the file as UTF-8.")
    except Exception as e:
        raise ValueError(f"Error processing TXT: {str(e)}")


def process_urls(urls: List[str]) -> List[LangChainDocument]:
    """
    Load and process content from URLs.
    
    Args:
        urls: List of URLs to process
        
    Returns:
        List of LangChain Document objects
        
    Raises:
        ValueError: If URL processing fails
    """
    try:
        # Filter out empty URLs
        valid_urls = [url.strip() for url in urls if url and url.strip()]
        
        if not valid_urls:
            raise ValueError("No valid URLs provided")
        
        documents = []
        errors = []
        
        for url in valid_urls:
            try:
                loader = WebBaseLoader(url)
                docs = loader.load()
                documents.extend(docs)
            except Exception as e:
                errors.append(f"Failed to load {url}: {str(e)}")
        
        if errors:
            # Log errors but continue if we got some documents
            for error in errors:
                logging.warning(error)
        
        if not documents:
            raise ValueError(
                "Failed to load any documents from the provided URLs. "
                "Please check the URLs and try again."
            )
        
        return documents
    
    except Exception as e:
        raise ValueError(f"Error processing URLs: {str(e)}")


def format_source_documents(source_docs: List[LangChainDocument], max_length: int = 500) -> str:
    """
    Format source documents for display.
    
    Args:
        source_docs: List of source documents
        max_length: Maximum length of each source preview
        
    Returns:
        Formatted string with source information
    """
    if not source_docs:
        return "No source documents available."
    
    formatted = []
    for i, doc in enumerate(source_docs, 1):
        content = doc.page_content
        preview = content[:max_length] + "..." if len(content) > max_length else content
        source = doc.metadata.get('source', 'Unknown')
        formatted.append(f"**Source {i}** ({source}):\n{preview}\n")
    
    return "\n".join(formatted)


def estimate_tokens(text: str) -> int:
    """
    Rough estimation of token count (approximately 4 characters per token).
    
    Args:
        text: Input text
        
    Returns:
        Estimated token count
    """
    return len(text) // 4


def sanitize_filename(filename: str) -> str:
    """
    Sanitize filename by removing invalid characters.
    
    Args:
        filename: Original filename
        
    Returns:
        Sanitized filename
    """
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename
