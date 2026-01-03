"""
Modern RAG Q&A Application
===========================
A production-ready Retrieval-Augmented Generation system using:
- Latest LangChain APIs (2024-2025)
- Modern Hugging Face integration
- FAISS vector store
- Streamlit UI

Author: Modernized Implementation
Date: January 2026
Python: 3.10+
"""


import streamlit as st
import logging
from typing import Optional, List, Union
from io import BytesIO

# Document processing imports (modern versions)
from docx import Document
from pypdf import PdfReader  # Modern replacement for PyPDF2
import numpy as np

# LangChain imports - using latest patterns
from langchain_community.document_loaders import WebBaseLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings, HuggingFaceEndpoint
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LangChainDocument

# Local imports
from config import Config
from utils import (
    validate_input,
    process_pdf,
    process_docx,
    process_txt,
    process_urls,
    setup_logging
)

# Initialize logging
logger = setup_logging()


class RAGSystem:
    """Encapsulates RAG system functionality with modern patterns."""
    
    def __init__(self, config: Config):
        """
        Initialize RAG system with configuration.
        
        Args:
            config: Configuration object with all settings
        """
        self.config = config
        self.embeddings = None
        self.llm = None
        self._initialize_models()
    
    def _initialize_models(self) -> None:
        """Initialize embedding and LLM models with error handling."""
        try:
            logger.info("Initializing embedding model...")
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.config.EMBEDDING_MODEL,
                model_kwargs={'device': self.config.DEVICE},
                encode_kwargs={'normalize_embeddings': True}  # Modern best practice
            )
            logger.info("Embedding model initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize embedding model: {e}")
            raise
    
    def _initialize_llm(self) -> None:
        """Initialize LLM lazily when needed."""
        if self.llm is None:
            try:
                logger.info(f"Initializing LLM: {self.config.LLM_MODEL} (Provider: {self.config.LLM_PROVIDER})")
                
                if self.config.LLM_PROVIDER == 'openai':
                    self.llm = ChatOpenAI(
                        model=self.config.LLM_MODEL,
                        temperature=self.config.TEMPERATURE,
                        max_tokens=self.config.MAX_NEW_TOKENS
                    )
                elif self.config.LLM_PROVIDER == 'huggingface':
                    self.llm = HuggingFaceEndpoint(
                        repo_id=self.config.LLM_MODEL,
                        huggingfacehub_api_token=self.config.HUGGINGFACE_API_KEY,
                        temperature=self.config.TEMPERATURE,
                        max_new_tokens=self.config.MAX_NEW_TOKENS,
                        timeout=120
                    )
                else:
                    raise ValueError(f"Unsupported LLM provider: {self.config.LLM_PROVIDER}")
                
                logger.info("LLM initialized successfully")
            except Exception as e:
                logger.error(f"Failed to initialize LLM: {e}")
                raise
    
    def process_input(
        self, 
        input_type: str, 
        input_data: Union[str, List[str], BytesIO, st.runtime.uploaded_file_manager.UploadedFile]
    ) -> Optional[FAISS]:
        """
        Process different input types and create a vector store.
        
        Args:
            input_type: Type of input (Link, PDF, Text, DOCX, TXT)
            input_data: The actual input data
            
        Returns:
            FAISS vector store or None on error
        """
        try:
            # Validate input
            if not validate_input(input_type, input_data):
                st.error("Invalid input data. Please check your input.")
                return None
            
            logger.info(f"Processing {input_type} input")
            
            # Process based on input type
            documents = []
            
            if input_type == "Link":
                documents = process_urls(input_data)
            elif input_type == "PDF":
                text = process_pdf(input_data)
                documents = [LangChainDocument(page_content=text, metadata={"source": "pdf"})]
            elif input_type == "Text":
                documents = [LangChainDocument(page_content=input_data, metadata={"source": "text"})]
            elif input_type == "DOCX":
                text = process_docx(input_data)
                documents = [LangChainDocument(page_content=text, metadata={"source": "docx"})]
            elif input_type == "TXT":
                text = process_txt(input_data)
                documents = [LangChainDocument(page_content=text, metadata={"source": "txt"})]
            else:
                st.error(f"Unsupported input type: {input_type}")
                return None
            
            if not documents:
                st.error("No content could be extracted from the input.")
                return None
            
            # Split documents using modern RecursiveCharacterTextSplitter
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.config.CHUNK_SIZE,
                chunk_overlap=self.config.CHUNK_OVERLAP,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]  # Better splitting
            )
            
            logger.info("Splitting documents into chunks...")
            splits = text_splitter.split_documents(documents)
            logger.info(f"Created {len(splits)} chunks")
            
            # Create vector store
            logger.info("Creating vector store...")
            with st.spinner("Creating vector embeddings..."):
                vector_store = FAISS.from_documents(
                    documents=splits,
                    embedding=self.embeddings
                )
            
            logger.info("Vector store created successfully")
            return vector_store
            
        except Exception as e:
            logger.error(f"Error processing input: {e}", exc_info=True)
            st.error(f"Error processing input: {str(e)}")
            return None
    
    def answer_question(self, vectorstore: FAISS, query: str) -> Optional[dict]:
        """
        Answer a question using the RAG pipeline.
        
        Args:
            vectorstore: FAISS vector store with documents
            query: User's question
            
        Returns:
            Dictionary with answer and source documents, or None on error
        """
        try:
            # Initialize LLM if not already done
            self._initialize_llm()
            
            logger.info(f"Processing query: {query[:100]}...")
            
            # Create retriever with modern configuration
            retriever = vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": self.config.TOP_K_RESULTS}
            )
            
            # Retrieve relevant documents
            docs = retriever.invoke(query)
            
            # Create context from documents
            context = "\n\n".join([doc.page_content for doc in docs])
            
            # Create prompt with context
            prompt = f"Answer the question based on the context:\n\nContext: {context}\n\nQuestion: {query}\n\nAnswer:"
            
            # Execute query
            with st.spinner("Generating answer..."):
                answer = self.llm.invoke(prompt)
                # Extract text based on provider
                if self.config.LLM_PROVIDER == 'openai':
                    answer_text = answer.content if hasattr(answer, 'content') else str(answer)
                else:  # huggingface returns string directly
                    answer_text = answer
                
                result = {
                    "result": answer_text,
                    "source_documents": docs
                }
            
            logger.info("Answer generated successfully")
            return result
            
        except Exception as e:
            logger.error(f"Error answering question: {e}", exc_info=True)
            st.error(f"Error generating answer: {str(e)}")
            return None


def main():
    """Main application entry point."""
    
    # Page configuration
    st.set_page_config(
        page_title="Modern RAG Q&A System",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Title and description
    st.title("🤖 Modern RAG Q&A System")
    st.markdown("""
    Upload documents, paste URLs, or enter text to create a knowledge base.
    Then ask questions and get AI-powered answers based on your content.
    """)
    
    # Load configuration
    try:
        config = Config()
    except Exception as e:
        st.error(f"Configuration error: {e}")
        st.info("Please ensure your .env file is set up correctly with HUGGINGFACE_API_KEY")
        st.stop()
    
    # Initialize RAG system (cached)
    if 'rag_system' not in st.session_state:
        try:
            with st.spinner("Initializing AI models..."):
                st.session_state.rag_system = RAGSystem(config)
        except Exception as e:
            st.error(f"Failed to initialize system: {e}")
            st.stop()
    
    rag_system = st.session_state.rag_system
    
    # Sidebar for settings
    with st.sidebar:
        st.header("⚙️ Settings")
        st.info(f"**Model:** {config.LLM_MODEL.split('/')[-1]}")
        st.info(f"**Embeddings:** {config.EMBEDDING_MODEL.split('/')[-1]}")
        
        # Option to clear vector store
        if st.button("🗑️ Clear Knowledge Base"):
            if 'vectorstore' in st.session_state:
                del st.session_state.vectorstore
                st.success("Knowledge base cleared!")
                st.rerun()
    
    # Main content area
    st.header("📚 Step 1: Add Knowledge")
    
    # Input type selection
    col1, col2 = st.columns([1, 3])
    with col1:
        input_type = st.selectbox(
            "Input Type",
            ["Text", "PDF", "DOCX", "TXT", "Link"],
            help="Select the type of content you want to add"
        )
    
    # Dynamic input based on type
    input_data = None
    
    if input_type == "Link":
        st.subheader("🔗 Web URLs")
        num_urls = st.number_input(
            "Number of URLs",
            min_value=1,
            max_value=20,
            value=1,
            step=1
        )
        input_data = []
        for i in range(int(num_urls)):
            url = st.text_input(f"URL {i+1}", key=f"url_{i}", placeholder="https://example.com")
            if url:
                input_data.append(url)
    
    elif input_type == "Text":
        st.subheader("📝 Text Input")
        input_data = st.text_area(
            "Enter your text",
            height=200,
            placeholder="Paste your text here..."
        )
    
    elif input_type == "PDF":
        st.subheader("📄 PDF Upload")
        input_data = st.file_uploader(
            "Upload PDF file",
            type=["pdf"],
            help="Upload a PDF document"
        )
    
    elif input_type == "TXT":
        st.subheader("📃 Text File Upload")
        input_data = st.file_uploader(
            "Upload text file",
            type=["txt"],
            help="Upload a plain text file"
        )
    
    elif input_type == "DOCX":
        st.subheader("📘 Word Document Upload")
        input_data = st.file_uploader(
            "Upload Word document",
            type=["docx", "doc"],
            help="Upload a Microsoft Word document"
        )
    
    # Process input button
    if st.button("🚀 Process & Create Knowledge Base", type="primary"):
        if not input_data or (isinstance(input_data, list) and not any(input_data)):
            st.warning("Please provide input data first.")
        else:
            vectorstore = rag_system.process_input(input_type, input_data)
            if vectorstore:
                st.session_state.vectorstore = vectorstore
                st.success("✅ Knowledge base created successfully!")
                st.balloons()
    
    # Q&A Section
    st.header("❓ Step 2: Ask Questions")
    
    if 'vectorstore' not in st.session_state:
        st.info("👆 Please create a knowledge base first by processing some content above.")
    else:
        st.success("✅ Knowledge base is ready!")
        
        # Query input
        query = st.text_input(
            "Ask your question",
            placeholder="What would you like to know?",
            help="Ask anything based on the documents you've uploaded"
        )
        
        # Submit button
        if st.button("🔍 Get Answer", type="primary"):
            if not query:
                st.warning("Please enter a question.")
            else:
                result = rag_system.answer_question(st.session_state.vectorstore, query)
                
                if result:
                    # Display answer
                    st.subheader("💡 Answer")
                    st.write(result['result'])
                    
                    # Display source documents if available
                    if 'source_documents' in result and result['source_documents']:
                        with st.expander("📚 View Source Documents"):
                            for i, doc in enumerate(result['source_documents'], 1):
                                st.markdown(f"**Source {i}:**")
                                st.text(doc.page_content[:500] + "..." if len(doc.page_content) > 500 else doc.page_content)
                                st.markdown("---")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: gray;'>
        <p>Modern RAG System | Built with LangChain, Hugging Face & Streamlit</p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
